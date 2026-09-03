# -*- coding: utf-8 -*-
"""
Builds "Udhëzues Përdorimi — Paneli Administrues" as a .docx.

The text lives in this file rather than in a separate Markdown source on purpose: every
paragraph is paired with the screenshot it explains, and keeping the two together is what
stops a renumbered screenshot from drifting away from its step.

Run:  python build_udhezuesi.py
Then: python to_pdf.py     (Word does the PDF, and refreshes the table of contents)
"""
import io
import os
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(HERE, 'screenshots')

BRAND = RGBColor(0x82, 0x36, 0x85)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x66, 0x66, 0x70)

VERSION = '1.0'
MONTHS_SQ = ['janar', 'shkurt', 'mars', 'prill', 'maj', 'qershor',
             'korrik', 'gusht', 'shtator', 'tetor', 'nëntor', 'dhjetor']


def sq_date(d):
    return '%d %s %d' % (d.day, MONTHS_SQ[d.month - 1], d.year)


# --------------------------------------------------------------------------- #
# low-level helpers
# --------------------------------------------------------------------------- #

def set_cell_background(cell, hex_colour):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def add_field(paragraph, instruction):
    """Insert a Word field (TOC, PAGE…). Word fills it in when the document opens."""
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = '…'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for node in (begin, instr, separate, placeholder, end):
        run._r.append(node)


def add_page_numbers(section):
    # The cover gets no number: a title page with a "1" on it reads as a draft. Word's
    # different-first-page switch leaves that footer empty and numbers everything after.
    section.different_first_page_header_footer = True
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, ' PAGE ')
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


# --------------------------------------------------------------------------- #
# document styling
# --------------------------------------------------------------------------- #

def build_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, colour, before, after in [
        ('Heading 1', 20, BRAND, 20, 10),
        ('Heading 2', 14, BRAND, 14, 6),
        ('Heading 3', 12, INK, 10, 4),
    ]:
        st = doc.styles[name]
        st.font.name = 'Calibri'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


# --------------------------------------------------------------------------- #
# content helpers
# --------------------------------------------------------------------------- #

class Guide:
    def __init__(self, doc):
        self.doc = doc
        self.figure = 0

    def h1(self, text):
        self.doc.add_page_break()
        self.doc.add_heading(text, level=1)

    def h2(self, text):
        self.doc.add_heading(text, level=2)

    def h3(self, text):
        self.doc.add_heading(text, level=3)

    def p(self, text, style=None):
        """`text` may use **bold** for UI labels — the only markup this needs."""
        para = self.doc.add_paragraph(style=style)
        for i, chunk in enumerate(text.split('**')):
            if chunk:
                run = para.add_run(chunk)
                run.bold = i % 2 == 1
        return para

    def steps(self, items):
        for item in items:
            self.p(item, style='List Number')

    def bullets(self, items):
        for item in items:
            self.p(item, style='List Bullet')

    def shot(self, filename, caption, width_cm=15.5):
        path = os.path.join(SHOTS, filename)
        if not os.path.exists(path):
            raise SystemExit('missing screenshot: %s' % filename)
        self.figure += 1
        self.doc.add_picture(path, width=Cm(width_cm))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run('Figura %d — %s' % (self.figure, caption))
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = MUTED
        cap.paragraph_format.space_after = Pt(14)

    def callout(self, title, body):
        """A single-cell shaded table — the "KUJDES" boxes."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, 'FDF3D6')

        head = cell.paragraphs[0]
        run = head.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(0x8A, 0x5A, 0x00)

        para = cell.add_paragraph()
        for i, chunk in enumerate(body.split('**')):
            if chunk:
                r = para.add_run(chunk)
                r.bold = i % 2 == 1
        self.doc.add_paragraph()

    def table(self, headers, rows, widths=None):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(cell, '823685')
        for row in rows:
            cells = t.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = ''
                para = cells[i].paragraphs[0]
                for j, chunk in enumerate(str(value).split('**')):
                    if chunk:
                        r = para.add_run(chunk)
                        r.bold = j % 2 == 1
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph()


# --------------------------------------------------------------------------- #
# the document
# --------------------------------------------------------------------------- #

def build():
    doc = Document()
    build_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    # ---------------- cover ----------------
    for _ in range(3):
        doc.add_paragraph()
    doc.add_picture(os.path.join(HERE, 'logo.png'), width=Cm(8.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(36)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Udhëzues Përdorimi')
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = BRAND

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Paneli Administrues')
    run.font.size = Pt(18)
    run.font.color.rgb = INK

    for _ in range(6):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Cacttus Education\n%s\nVersioni %s' % (sq_date(date.today()), VERSION))
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    # ---------------- table of contents ----------------
    doc.add_page_break()
    # Styled by hand rather than as Heading 1, so the contents page does not list itself.
    contents = doc.add_paragraph()
    run = contents.add_run('Përmbajtja')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = BRAND
    contents.paragraph_format.space_after = Pt(12)

    toc = doc.add_paragraph()
    add_field(toc, r' TOC \o "1-2" \h \z \u ')
    note = doc.add_paragraph()
    run = note.add_run('Nëse përmbajtja duket bosh, klikoni mbi të dhe shtypni F9 për ta rifreskuar.')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED

    add_page_numbers(section)

    g = Guide(doc)
    write_chapters(g)

    out = os.path.join(HERE, 'Udhezues-Perdorimi-Paneli-Administrues.docx')
    doc.save(out)
    print('  wrote %s (%d figures)' % (os.path.basename(out), g.figure))
    return out


def write_chapters(g):
    # =======================================================================
    g.h1('1. Hyrja')

    g.p('Ky udhëzues është për stafin që punon çdo ditë me panelin administrues të '
        'Cacttus Education: për ata që shtojnë trajnime, shkruajnë lajme dhe merren me '
        'aplikimet që vijnë nga faqja publike.')
    g.p('Nuk ju duhet asnjë njohuri teknike. Çdo kapitull ka hapa të numëruar dhe '
        'fotografi të ekranit, ashtu siç do ta shihni edhe ju.')

    g.h2('1.1 Si të kyçeni')
    g.steps([
        'Hapni shfletuesin dhe shkoni te https://admin.cacttus.education',
        'Shkruani email-in tuaj të punës dhe fjalëkalimin.',
        'Klikoni **Kyçu**.',
    ])
    g.shot('01-kycja.png', 'Faqja e kyçjes.', 12)
    g.p('Llogaritë i krijon administratori. Nëse nuk keni ende një llogari, kërkojini '
        'administratorit t’ju krijojë një.')

    g.h2('1.2 Çfarë shihni në Përmbledhje')
    g.p('Pas kyçjes hapet faqja **Përmbledhje**. Këtu shihni një pamje të shpejtë të '
        'gjendjes: sa aplikime kanë ardhur, sa trajnime dhe artikuj keni, dhe lëvizjet '
        'e fundit.')
    g.p('Në të majtë është menyja. Prej saj shkoni te **Aplikimet**, **Trajnimet**, '
        '**Format**, **Lajme**, **Regjistri**, **Përdoruesit** dhe **Cilësimet**.')
    g.shot('02-permbledhje.png', 'Faqja Përmbledhje dhe menyja në të majtë.')

    g.h2('1.3 Administrator apo Editor')
    g.p('Ka dy lloje llogarish. Roli juaj shkruhet lart djathtas, pranë emrit tuaj.')
    g.table(
        ['Veprimi', 'Administrator', 'Editor'],
        [
            ['Sheh dhe trajton aplikimet', 'Po', 'Po'],
            ['Krijon dhe ndryshon trajnime', 'Po', 'Po'],
            ['Krijon dhe ndryshon artikuj', 'Po', 'Po'],
            ['Krijon dhe ndryshon forma', 'Po', 'Po'],
            ['Shton dhe riemërton kategori', 'Po', 'Po'],
            ['Fshin një trajnim, artikull ose formë', 'Po', 'Jo'],
            ['Fshin një kategori', 'Po', 'Jo'],
            ['Krijon llogari për stafin', 'Po', 'Jo'],
        ],
        widths=[8.5, 3.5, 3.5],
    )
    g.p('Me pak fjalë: të dy rolet krijojnë dhe ndryshojnë përmbajtje. Vetëm '
        'administratori fshin.')

    # =======================================================================
    g.h1('2. Aplikimet')

    g.p('Te **Aplikimet** vijnë të gjitha kërkesat që dërgohen nga faqja publike: '
        'aplikimet për trajnime, mesazhet nga forma e kontaktit dhe kërkesat e bizneseve.')
    g.shot('03-aplikimet-lista.png', 'Lista e aplikimeve.')

    g.h2('2.1 Si të gjeni një aplikim')
    g.p('Mbi listë janë katër filtra. I përdorni sa herë lista bëhet e gjatë.')
    g.bullets([
        '**Kërko** — shkruani emrin, email-in ose numrin e telefonit.',
        '**Forma** — tregoni prej cilës formë ka ardhur aplikimi.',
        '**Tipi** — ndan aplikimet për trajnime nga mesazhet e kontaktit.',
        '**Statusi** — shfaq vetëm ato të reja, të kontaktuarat ose të arkivuarat.',
    ])
    g.shot('04-aplikimet-filtrat.png', 'Filtrat mbi listën e aplikimeve.')

    g.h2('2.2 Hapja e një aplikimi')
    g.p('Klikoni mbi emrin e aplikuesit. Hapet faqja me të gjitha të dhënat: kontaktet, '
        'forma prej së cilës ka ardhur dhe **Përgjigjet** — pyetjet shtesë që ka '
        'plotësuar aplikuesi.')
    g.shot('05-aplikimi-detajet.png', 'Faqja e një aplikimi.')
    g.p('Djathtas keni **Ndrysho statusin** dhe butonin **Dërgo email**, që hap '
        'programin tuaj të email-it me adresën e aplikuesit të plotësuar.')

    g.h2('2.3 Statuset')
    g.p('Çdo aplikim kalon nëpër tri gjendje:')
    g.table(
        ['Statusi', 'Kur përdoret'],
        [
            ['**I ri**', 'Sapo ka ardhur. Askush nuk e ka trajtuar ende.'],
            ['**I kontaktuar**', 'E keni marrë në telefon ose i keni shkruar.'],
            ['**I arkivuar**', 'Puna me të ka mbaruar, ose nuk ishte i vlefshëm.'],
        ],
        widths=[4.5, 11.0],
    )
    g.p('Nga lista, butoni **Shëno të kontaktuar** e ndryshon statusin me një klikim. '
        'Nëse e shënoni gabimisht, butoni **Kthe në «I ri»** e zhbën.')

    g.h2('2.4 Eksporti në Excel')
    g.steps([
        'Vendosni filtrat ashtu si i doni — eksportohet ajo që shihni.',
        'Klikoni **Eksporto CSV** lart djathtas.',
        'Ruani skedarin dhe hapeni me Excel.',
    ])
    g.callout(
        'Nëse shkronjat shqipe duken gabim në Excel',
        'Mos e hapni skedarin me dy klikime. Në Excel zgjidhni Data → From Text/CSV, '
        'zgjidhni skedarin dhe te "File Origin" vendosni **UTF-8**. Pastaj shkronjat ë '
        'dhe ç do të duken si duhet.')

    # =======================================================================
    g.h1('3. Trajnimet')

    g.p('Te **Trajnimet** menaxhoni katalogun që shihet në faqen publike. Çdo trajnim ka '
        'dy pjesë: kartelën (ajo që duket në listë) dhe faqen e vet.')
    g.shot('06-trajnimet-lista.png', 'Lista e trajnimeve.')

    g.h2('3.1 Krijimi i një trajnimi')
    g.steps([
        'Klikoni **Krijo trajnim të ri**.',
        'Shkruani **Titullin**. Linku publik gjenerohet vetë nga titulli.',
        'Zgjidhni **Kategorinë**.',
        'Zgjidhni **Formatin** dhe, nëse duhet, **Qytetin**.',
        'Plotësoni orët, çmimin dhe ligjëruesin.',
        'Poshtë, në pjesën e faqes, shtoni përshkrimin dhe pikat e forta.',
        'Klikoni **Ruaj**.',
    ])
    g.shot('07-trajnimi-kartela.png', 'Kartela e trajnimit.')

    g.h2('3.2 Kategoria dhe shtimi i një kategorie të re')
    g.p('Kategoria vendos se nën cilin filtër shfaqet trajnimi në faqen publike.')
    g.shot('08-trajnimi-kategoria.png', 'Fusha Kategoria me butonin Shto kategori.', 12)
    g.p('Nëse kategoria që ju duhet nuk ekziston, klikoni **+ Shto kategori**. Hapet një '
        'dritare e vogël: shkruani emrin dhe klikoni **Shto**. Kategoria e re zgjidhet '
        'menjëherë te trajnimi që po e plotësoni.')
    g.shot('11-trajnimi-dialogu-kategorive.png', 'Dritarja e kategorive.', 13)

    g.h2('3.3 Formati dhe qyteti')
    g.p('**Formati** tregon si zhvillohet trajnimi: në klasë, online ose hibrid.')
    g.callout(
        'Trajnimet online nuk kanë qytet',
        'Kur formati është **Online**, lëreni qytetin bosh. Trajnimi nuk mbahet askund '
        'fizikisht, ndaj në faqen publike nuk duhet të shfaqet asnjë qytet. Në filtrat e '
        'faqes ai gjendet te formati, jo te qyteti.')

    g.h2('3.4 Planprogrami në PDF')
    g.p('Te **Shkarko planprogramin (PDF)** ngarkoni dokumentin që vizitori mund ta '
        'shkarkojë nga faqja e trajnimit. Tërhiqeni skedarin ose klikoni për ta zgjedhur.')
    g.shot('10-trajnimi-planprogrami.png', 'Ngarkimi i planprogramit.', 13)

    g.h2('3.5 Lidhja me formën e aplikimit')
    g.p('Çdo trajnim duhet të tregojë se ku shkojnë aplikimet e tij. Te **Forma e '
        'aplikimit**, në fushën **Forma**, zgjidhni formën që pranon aplikimet për këtë '
        'trajnim. Butoni «Apliko» në faqen e trajnimit dërgon pikërisht atje.')
    g.shot('09-trajnimi-forma-aplikimit.png', 'Forma e aplikimit, radha dhe statusi.')

    g.h2('3.6 Aktivizimi dhe çaktivizimi')
    g.bullets([
        'Çelësi **Aktive** vendos nëse trajnimi duket fare në faqen publike. I fikur, '
        'ai zhduket nga lista.',
        '**Statusi** është diçka tjetër: **Aktive** do të thotë që ende pranohen '
        'aplikime, ndërsa **Përfunduar** e lë trajnimin të dukshëm, por të shënuar si '
        'të mbaruar.',
        '**Radha** vendos vendin në listë. Numri më i vogël shfaqet i pari.',
    ])

    g.h2('3.7 Çfarë ndodh pas ruajtjes')
    g.p('Sapo klikoni **Ruaj**, ndryshimi shkon te faqja publike. Deri në një minutë '
        'mund t’ju duhet të rifreskoni faqen për ta parë.')
    g.shot('18-web-trajnimet.png', 'Trajnimet siç i sheh vizitori.')

    # =======================================================================
    g.h1('4. Artikujt (Lajmet)')

    g.p('Te **Lajme** shkruani njoftimet dhe artikujt që dalin në faqen publike.')
    g.shot('12-lajmet-lista.png', 'Lista e artikujve.')

    g.h2('4.1 Krijimi dhe ndryshimi')
    g.steps([
        'Klikoni **Shto artikull**.',
        'Shkruani **Titullin**. Linku publik gjenerohet vetë nga titulli.',
        'Shkruani tekstin te **Përmbajtja**.',
        'Zgjidhni **Kategorinë** dhe ngarkoni **Foton e ballinës**.',
        'Kur teksti është gati, ndizni **Publikuar** dhe klikoni **Ruaj**.',
    ])
    g.shot('13-artikulli-editori.png', 'Editori i artikullit.')

    g.h2('4.2 Formatimi i tekstit')
    g.p('Mbi kutinë e tekstit është shiriti me butonat e formatimit. Zgjidhni tekstin '
        'dhe klikoni butonin që doni:')
    g.bullets([
        '**B** dhe **I** — të trasha dhe të pjerrëta.',
        '**H2** dhe **H3** — tituj brenda artikullit.',
        'Lista me pika dhe lista me numra.',
        'Thonjëza për një citat.',
        'Butoni i linkut dhe butoni i fotos.',
    ])
    g.p('Gjithçka tjetër hiqet automatikisht kur ruani. Prandaj mos ngjitni tekst me '
        'ngjyra e madhësi nga Word — ngjyrat nuk ruhen.')

    g.h2('4.3 Fotoja e ballinës')
    g.p('Djathtas, te **Fotoja e ballinës**, klikoni **Ngarko foto** ose tërhiqeni foton '
        'brenda kutisë. Pranohen **PNG, JPG, WEBP dhe GIF**, deri në **5 MB**.')

    g.h2('4.4 Kategoria dhe publikimi')
    g.shot('14-artikulli-kategoria.png', 'Publikimi dhe kategoria e artikullit.', 11)
    g.bullets([
        'Nëse kategoria ju mungon, klikoni **+ Shto kategori** — njësoj si te trajnimet.',
        'Kategoria nuk është e detyrueshme. Pa të, artikulli del vetëm te «Të gjitha».',
        'Çelësi **Publikuar** i fikur do të thotë draft: e ruani punën, por vizitori '
        'nuk e sheh.',
    ])

    g.h2('4.5 Si duket në faqen publike')
    g.p('Artikujt e publikuar dalin te /lajme. Sipër listës shfaqen kategoritë si '
        'butona: **Të gjitha** dhe një buton për secilën kategori që ka artikuj.')
    g.shot('19-web-lajmet.png', 'Lajmet siç i sheh vizitori.')

    # =======================================================================
    g.h1('5. Format')

    g.p('Një **formë** është grupi i pyetjeve që plotëson vizitori. Emri, email-i dhe '
        'telefoni mblidhen gjithmonë; te forma shtoni vetëm pyetjet shtesë.')
    g.shot('15-format-lista.png', 'Lista e formave.')

    g.h2('5.1 Krijimi i një forme për një trajnim')
    g.steps([
        'Te **Format**, klikoni **Krijo formë të re**.',
        'Shkruani **Titullin**, p.sh. «Aplikim — Web Development».',
        'Zgjidhni **Programin**.',
        'Klikoni **Shto fushë** për çdo pyetje shtesë.',
        'Klikoni **Ruaj formën**.',
        'Pastaj hapni trajnimin dhe zgjidhni këtë formë te **Forma e aplikimit**.',
    ])
    g.shot('16-forma-fushat.png', 'Të dhënat e formës dhe fusha e parë.')

    g.h2('5.2 Fushat')
    g.bullets([
        '**Pyetja** — teksti që lexon aplikuesi.',
        '**Lloji** — tekst i shkurtër, tekst i gjatë, listë zgjedhëse e kështu me radhë.',
        '**Çelësi i ruajtjes** — emri i shkurtër me të cilin ruhet përgjigjja. '
        'Plotësohet vetë; ndryshojeni vetëm nëse e dini pse.',
        '**E detyrueshme** — nëse aplikuesi nuk mund ta lërë bosh.',
    ])

    g.h2('5.3 Opsionet e një liste')
    g.p('Kur lloji është listë zgjedhëse, shtoni opsionet me **Shto opsion**. Çdo opsion '
        'ka dy kuti.')
    g.shot('17-forma-opsionet.png', 'Opsionet e një liste zgjedhëse.', 14)
    g.callout(
        'Shkruajini të njëjta',
        'Kutia e majtë është **Vlera e ruajtur** dhe e djathta **Teksti që sheh '
        'aplikuesi**. Shkruajini njësoj në të dyja. Nëse ndryshojnë, përgjigjja ruhet '
        'ndryshe nga ç’e lexoni ju në raport.')

    g.h2('5.4 Mbyllja e një forme')
    g.p('Çelësi **Pranon aplikime** vendos nëse forma pranon dërgesa. I fikur, faqja '
        'publike tregon se forma nuk është aktive dhe askush nuk mund të aplikojë.')
    g.callout(
        'KUJDES — tri forma nuk duhen fshirë kurrë',
        'Format **Kontakt**, **Kontakt biznesi** dhe **Rezervo klasë** janë të lidhura '
        'drejtpërdrejt me faqet e sitit: faqja e kontaktit, faqet për biznese dhe '
        'rezervimi i klasave. Nëse i fshini, ato faqe ndalojnë së pranuari mesazhe dhe '
        'vizitori merr gabim. Nëse doni t’i mbyllni përkohësisht, mos i fshini — thjesht '
        'fikni **Pranon aplikime**.')

    # =======================================================================
    g.h1('6. Kategoritë')

    g.p('Kategoritë janë etiketat që ndajnë përmbajtjen në faqen publike. Ka dy lista të '
        'ndara, sepse ato nuk përzihen:')
    g.bullets([
        'Kategoritë e **trajnimeve** — menaxhohen nga faqja e një trajnimi.',
        'Kategoritë e **artikujve** — menaxhohen nga faqja e një artikulli.',
    ])
    g.p('Në të dyja rastet klikoni **+ Shto kategori** pranë fushës **Kategoria**. Në '
        'dritaren që hapet mund të shtoni një kategori të re, ta riemërtoni një '
        'ekzistuese, ose — nëse jeni administrator — ta fshini.')
    g.p('Pranë çdo kategorie shkruhet sa trajnime ose artikuj e përdorin. Ai numër ju '
        'thotë paraprakisht nëse ajo mund të fshihet.')

    g.callout(
        'Pse nuk fshihet një kategori në përdorim',
        'Nëse provoni të fshini një kategori që përdoret, sistemi e ndalon dhe ju tregon '
        'sa trajnime ose artikuj varen prej saj. Kjo është me qëllim: përndryshe ata do '
        'të mbeteshin pa etiketë pa e vënë re askush. Zhvendosini së pari në një kategori '
        'tjetër, pastaj fshijeni.')

    # =======================================================================
    g.h1('7. Pyetje të shpeshta')

    g.h3('Harrova fjalëkalimin. Çfarë bëj?')
    g.p('Kontaktoni administratorin. Ai jua rivendos fjalëkalimin nga faqja '
        '**Përdoruesit**. Fjalëkalimin nuk e rikuperon dot vetë nga faqja e kyçjes.')

    g.h3('Fotoja nuk po ngarkohet')
    g.p('Kontrolloni dy gjëra:')
    g.bullets([
        'Formatin — pranohen vetëm **PNG, JPG, WEBP dhe GIF**. Skedarët **SVG nuk '
        'pranohen**, edhe pse duken si fotografi.',
        'Madhësinë — deri në **5 MB**. Nëse fotoja është më e madhe, zvogëlojeni para se '
        'ta ngarkoni.',
    ])

    g.h3('Në faqe shkruan se forma nuk pranon aplikime')
    g.p('Hapni **Format**, gjeni formën përkatëse dhe ndizni **Pranon aplikime**. Nëse '
        'edhe pas kësaj nuk punon, hapni trajnimin dhe sigurohuni që te **Forma e '
        'aplikimit** është zgjedhur forma e duhur.')

    g.h3('Ku shkojnë aplikimet?')
    g.p('Të gjitha vijnë te **Aplikimet**, pavarësisht se nga cila formë janë dërguar. '
        'Përdorni filtrin **Forma** për t’i ndarë. Askush nuk merr email automatik — '
        'aplikimet duhen parë në panel.')

    g.h3('Sa vonon të shfaqet një ndryshim në faqen publike?')
    g.p('Zakonisht menjëherë, por deri në rreth një minutë. Nëse nuk e shihni, '
        'rifreskoni faqen. Nëse as pas disa minutash nuk duket, kontrolloni që trajnimi '
        'të jetë **Aktive** ose artikulli **Publikuar**.')

    g.h3('E fshiva diçka pa dashje')
    g.p('Kontaktoni administratorin. Te **Regjistri** ruhet gjurma e çdo veprimi — kush '
        'e bëri dhe kur — dhe kjo ndihmon për ta rikthyer.')


if __name__ == '__main__':
    build()
